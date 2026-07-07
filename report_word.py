# -*- coding: utf-8 -*-
"""
Kimteks Tekstil_v1.docx formatinda Word raporu uretir.
Font: Arial | Header: #44546A | Grafik renkleri: #4472C4, #5B9BD5, #002060, #8895B6
"""
import os, re, tempfile, shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from nace_config import SECTOR_NAMES, SITC_MAP, get_kko_code, TOTAL_MANUFACTURING

# ─── RENK SABITLERI (dashboard ile ortak tasarim dili) ───────────────────────
CLR_HEADER_BG  = '1E3A5F'   # Derin lacivert — baslik arkaplan
CLR_WHITE      = 'FFFFFF'
CLR_BLUE_TITLE = '1D4ED8'   # Sekil basliklari — marka mavisi
CLR_BODY       = '0F1729'
CLR_CAPTION    = '64748B'   # Kaynak satirlari
CLR_POS        = '059669'   # Pozitif deger
CLR_NEG        = 'DC2626'   # Negatif deger
CLR_PANEL      = 'F1F5F9'   # Tablo panel dolgusu

ARIAL = 'Arial'

# ─── YARDIMCI: XML ───────────────────────────────────────────────────────────
def _set_cell_fill(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _set_para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def _para_spacing(para, before=0, after=80):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:before'), str(before))
    spc.set(qn('w:after'),  str(after))
    pPr.append(spc)

def _run(para, text, bold=False, italic=False, size_pt=10,
         color=CLR_BODY, font=ARIAL):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name   = font
    run.font.size   = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color)
    return run

def _heading_para(doc, text, level=1):
    """
    Kimteks tarzı mavi paragraf başlık (## ŞEKİL titresi değil,
    ana bölüm başlığı için).
    """
    para = doc.add_paragraph()
    _set_para_shading(para, CLR_HEADER_BG)
    _para_spacing(para, before=120, after=80)
    para.paragraph_format.left_indent  = Cm(0.3)
    para.paragraph_format.right_indent = Cm(0.3)
    _run(para, text, bold=True, size_pt=11, color=CLR_WHITE)
    return para

def _sekil_title(doc, text):
    """Şekil N: ... başlığı — örnek belge: kalın SİYAH, Arial 10pt, ortalı."""
    para = doc.add_paragraph()
    _para_spacing(para, before=120, after=60)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(para, text, bold=True, size_pt=10, color=CLR_BODY)
    return para

def _kaynak_line(doc, text):
    """Kaynak: ... satiri — örnek belge: Arial 8pt, ortalı, siyah (italik değil)."""
    para = doc.add_paragraph()
    _para_spacing(para, before=20, after=120)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(para, text, italic=False, size_pt=8, color=CLR_BODY)
    return para

def _yorum_para(doc, text):
    """Şekil yorumu — örnek belge: Arial 11pt, iki yana yaslı (justify), akıcı paragraf."""
    para = doc.add_paragraph()
    _para_spacing(para, before=0, after=120)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.right_indent = Cm(0.2)
    _run(para, text.strip(), size_pt=11, color=CLR_BODY)
    return para

def _blank(doc, h_pt=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run()
    run.font.size = Pt(h_pt)
    return para

def _add_image(doc, img_path, width_cm=16.0):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(para, before=20, after=20)
    run = para.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return para

# ─── OZET GOSTERGE TABLOSU ────────────────────────────────────────────────────
TR_MONTHS = ['Oca', 'Şub', 'Mar', 'Nis', 'May', 'Haz',
             'Tem', 'Ağu', 'Eyl', 'Eki', 'Kas', 'Ara']

def _tr_period(p):
    try:
        yr, mo = p.split('-')
        return f'{TR_MONTHS[int(mo)-1]} {yr}'
    except Exception:
        return p or '—'

def _last_val(fig_dict):
    """Ilk serinin son (donem, deger) çifti."""
    if not fig_dict:
        return None, None
    for _, series in fig_dict.items():
        pts = sorted(series.items())
        if pts:
            return pts[-1][1], pts[-1][0]
    return None, None

def _kpi_table(doc, fig1, fig3, fig4, fig5):
    """Rapor basina 'Ozet Gostergeler' tablosu (renkli degerler)."""
    prod_v, prod_p = _last_val(fig1)
    ih = {k: v for k, v in (fig3 or {}).items() if 'hracat' in k and 'thalat' not in k}
    it = {k: v for k, v in (fig3 or {}).items() if 'thalat' in k}
    ih_v, ih_p = _last_val(ih)
    it_v, it_p = _last_val(it)
    kko_sec = {k: v for k, v in (fig4 or {}).items() if 'anayii' not in k}
    kko_v, kko_p = _last_val({k: (dict(v) if not isinstance(v, dict) else v)
                              for k, v in kko_sec.items()})
    ufe_v, ufe_p = _last_val(fig5)

    rows = [
        ('Üretim (YoY)',  prod_v, prod_p, True),
        ('İhracat (YoY)', ih_v,   ih_p,   True),
        ('İthalat (YoY)', it_v,   it_p,   True),
        ('KKO',           kko_v,  kko_p,  False),
        ('ÜFE (Yıllık)',  ufe_v,  ufe_p,  True),
    ]
    rows = [r for r in rows if r[1] is not None]
    if not rows:
        return

    tbl = doc.add_table(rows=2, cols=len(rows))
    tbl.alignment = 1  # center
    for ci, (name, val, per, signed) in enumerate(rows):
        # Baslik hucresi
        c0 = tbl.cell(0, ci)
        c0.text = ''
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p0, name, bold=True, size_pt=7.5, color=CLR_WHITE)
        _set_cell_fill(c0, CLR_HEADER_BG)

        # Deger hucresi
        c1 = tbl.cell(1, ci)
        c1.text = ''
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if signed:
            vclr = CLR_POS if val >= 0 else CLR_NEG
            vtxt = f'{val:+.1f}%'
        else:
            vclr = CLR_HEADER_BG
            vtxt = f'{val:.1f}%'
        _run(p1, vtxt, bold=True, size_pt=12, color=vclr)
        p1b = c1.add_paragraph()
        p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p1b, _tr_period(per), size_pt=6.5, color=CLR_CAPTION)
        _set_cell_fill(c1, CLR_PANEL)
    return tbl

# ─── PRODTR ÜRÜN TABLOSU ─────────────────────────────────────────────────────
def _prodtr_table(doc, agg):
    """İlk 10 ürünü satış değerine göre tabloya döker (İSO tablosu ile aynı stil)."""
    top = (agg or {}).get('top') or []
    if not top:
        return
    yil = agg.get('yil', '')
    headers = ['Sıra', 'Ürün (PRODTR)', f'Satış Değeri (Milyar TL, {yil})', 'Girişim']
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = 1
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, size_pt=8, color=CLR_WHITE)
        _set_cell_fill(c, CLR_HEADER_BG)
    for i, t in enumerate(top[:10], 1):
        row = tbl.add_row().cells
        sv = (t.get('satis_deger') or 0) / 1e9
        vals = [str(i), t.get('tanim', '')[:70],
                f"{sv:,.2f}".replace(",", "\x00").replace(".", ",").replace("\x00", "."),
                str(t.get('girisim') or '—')]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER]
        for ci, (v, al) in enumerate(zip(vals, aligns)):
            c = row[ci]; c.text = ''
            p = c.paragraphs[0]; p.alignment = al
            _run(p, v, size_pt=8, color=CLR_BODY)
            if i % 2 == 0:
                _set_cell_fill(c, CLR_PANEL)
    _table_borders(tbl)
    return tbl

# ─── İSO FIRMA TABLOSU (örnek belge formatı) ─────────────────────────────────
def _iso_table(doc, iso_agg):
    """İlk 10 kuruluşu örnek belgedeki tablo düzeninde ekler."""
    top = (iso_agg or {}).get('top10') or []
    if not top:
        return
    yil = iso_agg.get('yil', '')
    headers = ['İSO Genel Sıra', 'Sektör Sıra', 'Kuruluş Adı',
               f'Üretimden Satışlar (Bin TL, {yil})']
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = 1
    tbl.autofit = True

    # Başlık satırı
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, size_pt=8, color=CLR_WHITE)
        _set_cell_fill(c, CLR_HEADER_BG)

    # Satırlar
    for i, t in enumerate(top, 1):
        row = tbl.add_row().cells
        vals = [
            str(t.get('sira') or '—'),
            str(i),
            str(t.get('firma') or ''),
            f"{(t.get('uretim_satis') or 0)/1000:,.0f}".replace(',', '.'),
        ]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]
        for ci, (v, al) in enumerate(zip(vals, aligns)):
            c = row[ci]; c.text = ''
            p = c.paragraphs[0]; p.alignment = al
            _run(p, v, size_pt=8, color=CLR_BODY)
            if i % 2 == 0:
                _set_cell_fill(c, CLR_PANEL)

    # İnce kenarlıklar
    _table_borders(tbl)
    return tbl

def _table_borders(tbl, color='BFBFBF', sz='4'):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)

# ─── LLM ANALIZ BLOGU ─────────────────────────────────────────────────────────
def parse_llm_sections(text):
    """
    LLM ciktisini [TAG] bolumlerine ayirir.
    Doner: {'GIRIS': '...', 'SEKIL1': '- ...\n- ...', ...}
    """
    result = {}
    current_key = 'GIRIS'
    current_lines = []
    for line in text.splitlines():
        m = re.match(r'^\[([A-Z0-9]+)\]\s*$', line.strip())
        if m:
            if current_lines:
                result[current_key] = '\n'.join(current_lines).strip()
            current_key   = m.group(1)
            current_lines = []
        else:
            current_lines.append(line)
    if current_lines:
        result[current_key] = '\n'.join(current_lines).strip()
    return result

def bullets_from_text(text):
    """Her satiri madde olarak listele (bos satirlari atla)."""
    lines = []
    for ln in text.splitlines():
        ln = ln.strip().lstrip('-–*• ').strip()
        if ln:
            lines.append(ln)
    return lines

# ─── ANA RAPOR URETICI ────────────────────────────────────────────────────────
def build_word_report(nace, fig1, fig2, fig3, fig4, fig5,
                      analysis_text, chart_paths, out_dir=None, iso_agg=None,
                      news_analysis=None, prodtr_agg=None):
    """
    Kimteks formatiyla Word raporu uretir.
    analysis_text: LLM'den gelen [TAG] bolumlu Turkce metin
    chart_paths: {'sekil1': '/tmp/s1.png', ...}
    """
    sector = SECTOR_NAMES.get(nace, nace)
    try:
        from report_charts import nace_name as _nn
        official = _nn(nace)
        if official and official != nace:
            sector = official  # TUIK codelist'indeki resmi Turkce isim
    except Exception:
        pass
    if out_dir is None:
        out_dir = os.path.expanduser("~\\Desktop")

    safe = re.sub(r'[^0-9A-Za-zÇĞİÖŞÜçğıöşü]+', '_', sector[:36]).strip('_')
    out_file = os.path.join(out_dir, f"{nace}_{safe}_Rapor.docx")

    doc = Document()

    # ─── Sayfa yapisi: A4, 2.54 cm (1 inç) kenar boşluk — örnek belgeyle aynı ──
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)

    # Varsayilan paragraf fontu
    style = doc.styles['Normal']
    style.font.name = ARIAL
    style.font.size = Pt(10)

    # ─── ANA BASLIK (örnek belge: ortalı "Sektörel Araştırmalar" + sektör adı) ──
    t1 = doc.add_paragraph()
    _para_spacing(t1, before=0, after=40)
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t1, 'Sektörel Araştırmalar', bold=True, size_pt=13, color=CLR_BODY)

    t2 = doc.add_paragraph()
    _para_spacing(t2, before=0, after=160)
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t2, f'{sector} Sektörü', bold=True, size_pt=12, color=CLR_BODY)

    # ─── LLM Analiz metnini parse et ─────────────────────────────────────────
    if analysis_text:
        sections = parse_llm_sections(analysis_text)
    else:
        sections = {}

    # GİRİŞ paragrafı (iki yana yaslı, akıcı)
    giris = sections.get('GIRIS', '').strip()
    if giris:
        for para_text in giris.split('\n'):
            para_text = para_text.strip()
            if para_text:
                _yorum_para(doc, para_text)

    # ─── ŞEKİL BLOKLARI ───────────────────────────────────────────────────────
    is_toplam = (nace == TOTAL_MANUFACTURING)
    sekil2_title  = ('İmalat Sanayii Alt Sektörleri (Önceki Yıla Göre Değişim, %)' if is_toplam
                    else f'{sector} Alt Kırılımlar (Önceki Yıla Göre Değişim, %)')
    sekil2_kaynak = ('Kaynak: TÜİK – Sanayi Üretim Endeksi, Ana Endeksler (2 Haneli NACE, C10–C33)' if is_toplam
                     else 'Kaynak: TÜİK – Sanayi Üretim Endeksi, Sınıf Düzeyi')

    sekil_defs = [
        ('sekil1', 'Şekil 1', f'{sector} Sanayi Üretim Endeksi (Önceki Yıla Göre Değişim, %)',
         'Kaynak: Türkiye İstatistik Kurumu (TÜİK)'),
        ('sekil2', 'Şekil 2', sekil2_title, 'Kaynak: TÜİK'),
        ('sekil3', 'Şekil 3', f'{sector} Sektörü Dış Ticaret Verileri (Önceki Yıla Göre Değişim, %)',
         'Kaynak: TÜİK'),
        ('sekil4', 'Şekil 4', 'Kapasite Kullanım Oranı (Yıllıklandırılmış, %)',
         'Kaynak: Türkiye Cumhuriyet Merkez Bankası (TCMB)'),
        ('sekil5', 'Şekil 5', 'Üretici Fiyat Endeksi Değişimi – Yıllık %',
         'Kaynak: TÜİK'),
    ]
    fig_data = [fig1, fig2, fig3, fig4, fig5]
    sekil_keys = ['SEKIL1', 'SEKIL2', 'SEKIL3', 'SEKIL4', 'SEKIL5']

    for (key, s_no, title, kaynak), fig_d, skey in zip(sekil_defs, fig_data, sekil_keys):
        if not fig_d and not chart_paths.get(key):
            continue

        # Sekil baslik (kalın siyah, ortalı)
        _sekil_title(doc, f'{s_no}: {title}')

        # Grafik
        img = chart_paths.get(key)
        if img and os.path.exists(img):
            _add_image(doc, img, width_cm=15.5)

        # Kaynak (ortalı, 8pt)
        _kaynak_line(doc, kaynak)

        # Yorum — akıcı paragraf (iki yana yaslı)
        yorum = sections.get(skey, '').strip()
        if yorum:
            for para_text in yorum.split('\n'):
                if para_text.strip():
                    _yorum_para(doc, para_text)

    # ─── SEKTÖRÜN KURUMSAL GÖRÜNÜMÜ (İSO 500) ────────────────────────────────
    if iso_agg:
        yil = iso_agg.get('yil', '')
        # Bölüm başlığı (örnek belgedeki "Firmanın Sektördeki Konumu" muadili)
        bh = doc.add_paragraph()
        _para_spacing(bh, before=160, after=100)
        bh.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(bh, 'Sektörün İSO 500 İçindeki Konumu', bold=True, size_pt=11, color=CLR_BODY)

        # Yorum paragrafı
        iso_llm = sections.get('SEKIL6', '').strip()
        if iso_llm:
            for pt in iso_llm.split('\n'):
                if pt.strip(): _yorum_para(doc, pt)
        else:
            _yorum_para(doc,
                f"Türkiye'nin en büyük sanayi kuruluşlarını kapsayan İSO 500 ve İkinci 500 "
                f"listelerinde {sector} sektöründen {yil} yılında {iso_agg['firma_sayisi']} "
                f"kuruluş yer almıştır (İSO 500: {iso_agg['firma_500']}, İkinci 500: "
                f"{iso_agg['firma_2_500']}). Bu kuruluşların toplam üretimden satışları "
                f"{iso_agg['toplam_uretim_satis']/1e9:.1f} milyar TL, toplam ihracatı ise "
                f"{iso_agg['toplam_ihracat_musd']:.0f} milyon dolar düzeyinde gerçekleşmiştir." +
                (f" Sektörün medyan FAVÖK marjı %{iso_agg['favok_marj_med']:.1f} olarak "
                 f"hesaplanmıştır." if iso_agg.get('favok_marj_med') is not None else ""))

        # Grafik (İlk 10 kuruluş)
        if chart_paths.get('sekil6'):
            _sekil_title(doc, f"Şekil 6: {sector} Sektöründe İlk 10 Kuruluş "
                              f"(Üretimden Satışlar, Milyar TL, {yil})")
            _add_image(doc, chart_paths['sekil6'], width_cm=15.5)
            _kaynak_line(doc, 'Kaynak: İstanbul Sanayi Odası (İSO 500 ve İkinci 500)')

        # Firma tablosu (örnek belgedeki İSO tablosu muadili)
        _iso_table(doc, iso_agg)
        _blank(doc, 6)

    # ─── ÜRÜN DETAYI (PRODTR) ────────────────────────────────────────────────
    if prodtr_agg and prodtr_agg.get('top'):
        yilp = prodtr_agg.get('yil', '')
        bh = doc.add_paragraph()
        _para_spacing(bh, before=160, after=100)
        _run(bh, 'Sektörün Ürün Bazında Görünümü', bold=True, size_pt=11, color=CLR_BODY)
        def _tl(v):
            return f"{v:,.1f}".replace(",", "\x00").replace(".", ",").replace("\x00", ".")
        _yorum_para(doc,
            f"TÜİK Sanayi Ürün İstatistikleri (PRODTR) verilerine göre sektörde "
            f"{prodtr_agg['urun_sayisi']} farklı ürün çeşidi tanımlı olup, {yilp} yılında "
            f"beyan edilen toplam ürün satış değeri {_tl(prodtr_agg['toplam_satis']/1e9)} milyar "
            f"TL düzeyindedir. En yüksek satış değerine sahip ürün "
            f"“{prodtr_agg['top'][0]['tanim']}” olup "
            f"{_tl((prodtr_agg['top'][0]['satis_deger'] or 0)/1e9)} milyar TL ile öne çıkmaktadır.")
        _prodtr_table(doc, prodtr_agg)
        _blank(doc, 6)

    # ─── SEKTÖREL BEKLENTİLER VE RİSK ANALİZİ (haber sentezi) ────────────────
    if news_analysis and news_analysis.strip():
        _blank(doc, 6)
        bh = doc.add_paragraph()
        _para_spacing(bh, before=120, after=100)
        _run(bh, 'Sektörel Beklentiler ve Risk Analizi', bold=True, size_pt=11, color=CLR_BODY)
        for b in bullets_from_text(news_analysis):
            _yorum_para(doc, b)
        _kaynak_line(doc,
            f'Kaynak: Güncel haber taraması ve yapay zekâ sentezi ('
            f'{datetime.now().strftime("%d.%m.%Y")}). Görüşler yatırım tavsiyesi değildir.')
        _blank(doc, 4)

    # ─── RAPOR ALT BİLGİSİ ───────────────────────────────────────────────────
    _blank(doc, 4)
    footer_p = doc.add_paragraph()
    _para_spacing(footer_p, before=60, after=0)
    _run(footer_p,
         f'Rapor tarihi: {datetime.now().strftime("%d.%m.%Y")}  ·  '
         f'Kaynak: TÜİK SDMX REST API v1.5 + TCMB EVDS  ·  '
         f'Otomatik üretilmiştir',
         size_pt=7, color=CLR_CAPTION, italic=True)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_file)
    return out_file
